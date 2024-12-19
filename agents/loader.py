import os
from langchain.schema import Document
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

import chardet  # Biblioteca para detectar a codificação

class DocumentLoader:
    """Agente para carregar e processar documentos."""

    @staticmethod
    def detect_encoding(file_path):
        """Detecta a codificação de um arquivo."""
        with open(file_path, "rb") as f:
            raw_data = f.read(1024)  # Lê os primeiros 1024 bytes para análise
            result = chardet.detect(raw_data)
            return result["encoding"]

    @staticmethod
    def extract_text(file_path):
        """Extrai texto de documentos suportados."""
        if file_path.endswith(".txt"):
            # Detecta a codificação automaticamente
            encoding = DocumentLoader.detect_encoding(file_path)
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        elif file_path.endswith(".pdf"):
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(file_path)
            return "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        elif file_path.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            raise ValueError(f"Formato de arquivo não suportado: {file_path}")

    @staticmethod
    def load_documents(folder_path):
        """Carrega todos os documentos de um diretório."""
        from langchain.schema import Document
        import os

        documents = []
        for file_name in os.listdir(folder_path):
            file_path = os.path.join(folder_path, file_name)
            try:
                content = DocumentLoader.extract_text(file_path)
                if content.strip():
                    documents.append(Document(page_content=content, metadata={"source": file_name}))
            except Exception as e:
                print(f"Erro ao processar '{file_name}': {e}")
        if not documents:
            raise ValueError("Nenhum documento válido encontrado.")
        return documents
